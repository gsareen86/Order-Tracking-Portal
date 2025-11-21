import pandas as pd
import random
from datetime import datetime, timedelta
from docx import Document
import os

# Constants
ORDER_STATUSES = ['Ordered', 'Packaging', 'Shipped', 'Delivered']
ORDER_TYPES = ['Packaging Films', 'Aseptic Liquid Packaging', 'Chemicals', 'Holography', 'Engineering Machinery']
BUYERS = [
    {
        "name": "XYZ Industries Pvt Ltd",
        "address": "Plot No 12, Block-A, Sector-XX, Gurgaon, Haryana",
        "gst": "GSTIN123456567"
    },
    {
        "name": "ABC Foods Corp",
        "address": "Industrial Area, Phase 2, Manesar, Haryana",
        "gst": "GSTIN987654321"
    },
    {
        "name": "Global Beverages Ltd",
        "address": "Tech Park, Whitefield, Bangalore, Karnataka",
        "gst": "GSTIN456789123"
    }
]
SELLER = {
    "name": "ABC Packaging Pvt Ltd",
    "address": "C-Block, 12 Street, Sector-62, Noida, UP",
    "tin": "TIN5676780000"
}

ITEMS = {
    'Packaging Films': ['BOPP Film', 'BOPET Film', 'CPP Film', 'Metalized Film'],
    'Aseptic Liquid Packaging': ['Aseptic Brick Pack', 'Aseptic Pillow Pack'],
    'Chemicals': ['Speciality Coating - Primer GD-II', 'Flexo Ink', 'Lamination Adhesive'],
    'Holography': ['Holographic Film', 'Security Labels', 'Hot Stamping Foil'],
    'Engineering Machinery': ['Slitting Machine', 'Pouch Making Machine', 'Printing Machine']
}

def generate_orders(num_orders=50):
    orders = []
    for i in range(num_orders):
        order_type = random.choice(ORDER_TYPES)
        item = random.choice(ITEMS[order_type])
        buyer = random.choice(BUYERS)
        
        order_date = datetime.now() - timedelta(days=random.randint(1, 180))
        status = random.choice(ORDER_STATUSES)
        
        # Logic for status and dates
        if status == 'Ordered':
            delivery_days = random.randint(5, 15)
            expected_delivery = order_date + timedelta(days=delivery_days)
        elif status == 'Packaging':
            delivery_days = random.randint(4, 12)
            expected_delivery = order_date + timedelta(days=delivery_days)
        elif status == 'Shipped':
            delivery_days = random.randint(2, 8)
            expected_delivery = order_date + timedelta(days=delivery_days)
        else: # Delivered
            delivery_days = random.randint(2, 10)
            expected_delivery = order_date + timedelta(days=delivery_days)
            if expected_delivery > datetime.now():
                expected_delivery = datetime.now() - timedelta(days=random.randint(0, 5))

        quantity = random.randint(10, 500) * 10
        unit_cost = random.randint(100, 5000)
        total_cost = quantity * unit_cost
        
        order = {
            "Order No": f"ORD-{10000+i}",
            "Order Date": order_date.strftime("%Y-%m-%d"),
            "Order Status": status,
            "Order Type": order_type,
            "Item": item,
            "Quantity": quantity,
            "Unit Cost": unit_cost,
            "Total Amount": total_cost,
            "Expected Delivery": expected_delivery.strftime("%Y-%m-%d %H:%M"),
            "Buyer Name": buyer["name"],
            "Buyer Address": buyer["address"],
            "Buyer GST": buyer["gst"],
            "Seller Name": SELLER["name"],
            "Seller Address": SELLER["address"],
            "Seller TIN": SELLER["tin"]
        }
        orders.append(order)
    
    df = pd.DataFrame(orders)
    os.makedirs('data', exist_ok=True)
    df.to_excel('data/order_db.xlsx', index=False)
    print("Generated data/order_db.xlsx")

def create_invoice_template():
    doc = Document()
    doc.add_heading('INVOICE', 0)
    
    # Seller Details
    doc.add_paragraph(f"Seller: {SELLER['name']}")
    doc.add_paragraph(f"Address: {SELLER['address']}")
    doc.add_paragraph(f"TIN: {SELLER['tin']}")
    
    doc.add_paragraph("-" * 50)
    
    # Buyer Details Placeholders
    doc.add_paragraph("Buyer: {{buyer_name}}")
    doc.add_paragraph("Address: {{buyer_address}}")
    doc.add_paragraph("GST No: {{buyer_gst}}")
    
    doc.add_paragraph("-" * 50)
    
    # Order Details
    doc.add_paragraph("Invoice No: {{invoice_no}}")
    doc.add_paragraph("Order Date: {{order_date}}")
    doc.add_paragraph("Order No: {{order_no}}")
    
    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Quantity'
    hdr_cells[2].text = 'Unit Cost (INR)'
    hdr_cells[3].text = 'Total Cost (INR)'
    
    row_cells = table.add_row().cells
    row_cells[0].text = "{{item_name}}"
    row_cells[1].text = "{{quantity}}"
    row_cells[2].text = "{{unit_cost}}"
    row_cells[3].text = "{{total_cost}}"
    
    doc.add_paragraph("\n")
    doc.add_paragraph("Total Amount: INR {{total_cost}}", style='Quote')
    
    doc.add_paragraph("\n\nAuthorized Signatory")
    
    os.makedirs('data', exist_ok=True)
    doc.save('data/invoice_template.docx')
    print("Generated data/invoice_template.docx")

if __name__ == "__main__":
    generate_orders()
    create_invoice_template()
