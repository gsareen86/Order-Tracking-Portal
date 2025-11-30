import pandas as pd
from docx import Document
from docx.shared import Pt
import os
import json
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from datetime import datetime

DATA_DIR = 'data'
DATA_DIR = 'data'
ORDER_DB_PATH = os.path.join(DATA_DIR, 'order_db_v2.xlsx')
INVOICE_TEMPLATE_PATH = os.path.join(DATA_DIR, 'invoice_template.docx')

def get_orders_df():
    if not os.path.exists(ORDER_DB_PATH):
        return pd.DataFrame()
    
    df = pd.read_excel(ORDER_DB_PATH)
    # Ensure consistency in column names (strip whitespace)
    df.columns = [c.strip() for c in df.columns]
    
    # Convert dates to string for JSON serialization
    if 'Order Date' in df.columns:
        df['Order Date'] = pd.to_datetime(df['Order Date'])
    if 'Expected Delivery' in df.columns:
        df['Expected Delivery'] = pd.to_datetime(df['Expected Delivery'])
    if 'Shipped Date' in df.columns:
         df['Shipped Date'] = pd.to_datetime(df['Shipped Date'])
    if 'Delivered Date' in df.columns:
         df['Delivered Date'] = pd.to_datetime(df['Delivered Date'])
    if 'Payment Due Date' in df.columns:
         df['Payment Due Date'] = pd.to_datetime(df['Payment Due Date'])
         
    return df

def get_all_orders():
    df = get_orders_df()
    if df.empty:
        return []
        
    # Convert timestamps to strings for JSON serialization
    for col in df.select_dtypes(include=['datetime64']).columns:
        df[col] = df[col].astype(str).replace('NaT', None)
        
    return df.to_dict('records')

def load_config():
    config_path = 'config.json'
    if os.path.exists(config_path):
        with open(config_path, 'r') as f:
            return json.load(f)
    return {"payment_due_days": 60}

def get_order_by_id(order_id):
    orders = get_all_orders()
    for order in orders:
        if order['Order No'] == order_id:
            return order
    return None

def cancel_order(order_id, reason):
    if not os.path.exists(ORDER_DB_PATH):
        return False
    
    df = pd.read_excel(ORDER_DB_PATH)
    if order_id in df['Order No'].values:
        df.loc[df['Order No'] == order_id, 'Order Status'] = 'Cancelled'
        # In a real app, we would store the cancellation reason somewhere
        df.to_excel(ORDER_DB_PATH, index=False)
        return True
    return False

def generate_invoice_docx(order_id):
    order = get_order_by_id(order_id)
    if not order:
        return None
    
    if not os.path.exists(INVOICE_TEMPLATE_PATH):
        return None

    doc = Document(INVOICE_TEMPLATE_PATH)
    
    replacements = {
        '{{buyer_name}}': str(order.get('Buyer Name', '')),
        '{{buyer_address}}': str(order.get('Buyer Address', '')),
        '{{buyer_gst}}': str(order.get('Buyer GST', '')),
        '{{invoice_no}}': f"INV-{order_id.split('-')[1]}",
        '{{order_date}}': str(order.get('Order Date', '')),
        '{{order_no}}': str(order.get('Order No', '')),
        '{{item_name}}': str(order.get('Item', '')),
        '{{quantity}}': str(order.get('Quantity', '')),
        '{{unit_cost}}': str(order.get('Unit Cost', '')),
        '{{total_cost}}': str(order.get('Total Amount', ''))
    }

    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    
    output_path = os.path.join(DATA_DIR, f'invoice_{order_id}.docx')
    doc.save(output_path)
    return output_path

def generate_invoice_pdf(order_id):
    order = get_order_by_id(order_id)
    if not order:
        return None
        
    output_path = os.path.join(DATA_DIR, f'invoice_{order_id}.pdf')
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()
    
    # Header
    elements.append(Paragraph(f"INVOICE", styles['Title']))
    elements.append(Spacer(1, 12))
    
    # Seller Info
    elements.append(Paragraph(f"<b>Seller:</b> {order.get('Seller Name', '')}", styles['Normal']))
    elements.append(Paragraph(f"<b>Address:</b> {order.get('Seller Address', '')}", styles['Normal']))
    elements.append(Paragraph(f"<b>TIN:</b> {order.get('Seller TIN', '')}", styles['Normal']))
    elements.append(Spacer(1, 12))
    
    # Buyer Info
    elements.append(Paragraph(f"<b>Buyer:</b> {order.get('Buyer Name', '')}", styles['Normal']))
    elements.append(Paragraph(f"<b>Address:</b> {order.get('Buyer Address', '')}", styles['Normal']))
    elements.append(Paragraph(f"<b>GST:</b> {order.get('Buyer GST', '')}", styles['Normal']))
    elements.append(Spacer(1, 12))
    
    # Order Info
    elements.append(Paragraph(f"<b>Invoice No:</b> INV-{order_id.split('-')[1]}", styles['Normal']))
    elements.append(Paragraph(f"<b>Order No:</b> {order_id}", styles['Normal']))
    elements.append(Paragraph(f"<b>Date:</b> {order.get('Order Date', '')}", styles['Normal']))
    elements.append(Spacer(1, 20))
    
    # Table Data
    data = [
        ['Item', 'Quantity', 'Unit Cost', 'Total Cost'],
        [order.get('Item', ''), str(order.get('Quantity', '')), f"INR {order.get('Unit Cost', '')}", f"INR {order.get('Total Amount', '')}"]
    ]
    
    # Add Advance and Balance if available
    total_amount = order.get('Total Amount', 0)
    advance_amount = order.get('Advance Amount', 0)
    balance_due = total_amount - advance_amount
    
    data.append(['', '', 'Advance Paid:', f"INR {advance_amount}"])
    data.append(['', '', 'Balance Due:', f"INR {balance_due}"])
    
    t = Table(data)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(t)
    
    elements.append(Spacer(1, 20))
    elements.append(Paragraph("Authorized Signatory", styles['Normal']))
    
    doc.build(elements)
    return output_path
